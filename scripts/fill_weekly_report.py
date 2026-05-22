from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

BULLET = "\u2022 "  # •

src = r"c:\Users\smits\Downloads\WEEKLY REPORT-1.docx"
out = r"c:\Users\smits\Downloads\WEEKLY REPORT-1-filled-v2.docx"

doc = Document(src)


def bullets(items):
    return "\n".join(f"{BULLET}{item}" for item in items)


def set_para(index, heading, items=None):
    p = doc.paragraphs[index]
    p.clear()
    if items:
        text = (heading + "\n" + bullets(items)) if heading else bullets(items)
    else:
        text = heading
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


# Personal fields
set_para(1, "College ID: [Enter College ID]")
set_para(2, "Project Title: Personalized AI Tutor (personalized-ai-tutor)")
set_para(18, "Student Id: [Enter Student ID]")
set_para(19, "Student Name: [Enter Full Name]")

# Work done in last week
set_para(
    5,
    "1. Project Setup & Architecture",
    [
        "Created project structure: Frontend (React), Backend (Express), Database (PostgreSQL).",
        "Configured Vite for frontend development and hot reload.",
        "Set up Express server with CORS and JSON middleware on port 5000.",
        "Established PostgreSQL connection pool using environment variables.",
    ],
)

set_para(
    6,
    "2. Database Design & Content",
    [
        "Created schema for users, topics, questions, quiz_results, progress, and recommendations.",
        "Defined primary keys, foreign keys, and timestamp fields.",
        "Seeded MCQ banks for 9 subjects: DSA, DBMS, OOP, OS, CN, Web Dev, Algorithms, AI/ML, DS.",
    ],
)

set_para(
    7,
    "3. Authentication & User Interface",
    [
        "Built Register and Login REST APIs with field validation.",
        "Hashed passwords using bcrypt before storing in database.",
        "Generated JWT token on login (valid for 7 days).",
        "Developed Landing, Sign Up, and Login pages with React Router.",
        "Styled UI using Tailwind CSS with responsive layout.",
        "Added marketing sections on landing page: hero, features, and CTA.",
    ],
)

set_para(
    8,
    "4. Quiz Module & Student Dashboard",
    [
        "Topic selection page showing subject name, description, and question count.",
        "Quiz start API returns 8 random MCQs per topic (without correct answers).",
        "Quiz submit API calculates score, percentage, and saves result to database.",
        "Post-quiz review shows explanation for each question.",
        "Dashboard API returns quizzes taken, weekly count, average score, and topics completed.",
        "Implemented routes: /, /signup, /login, /dashboard, /topics, /quiz/:topicId.",
    ],
)

# Reason for incomplete work
set_para(
    10,
    "",
    [
        "AI-based personalized tutoring is not implemented yet.",
        "Recommendation system table exists but logic is not built.",
        "JWT middleware is not connected to protected API routes.",
        "Learning streak feature is pending (placeholder value used).",
        "Adaptive difficulty-based questions are not developed.",
        "Full system testing and deployment are still pending.",
    ],
)

# Plans for next week
set_para(
    12,
    "",
    [
        "Secure dashboard and quiz APIs using JWT authentication.",
        "Start AI tutor module for weak-topic detection and recommendations.",
        "Link progress table to track learning status per topic.",
        "Add quiz history and detailed answer review on frontend.",
        "Implement daily learning streak on dashboard.",
        "Run integration testing and fix remaining bugs.",
        "Write technical documentation and setup deployment.",
    ],
)

# References
set_para(
    15,
    "",
    [
        "React — https://react.dev/",
        "Express.js — https://expressjs.com/",
        "PostgreSQL — https://www.postgresql.org/docs/",
        "JWT — https://jwt.io/introduction",
        "Tailwind CSS — https://tailwindcss.com/docs",
        "bcrypt (npm) — https://www.npmjs.com/package/bcrypt",
        "Vite — https://vite.dev/",
    ],
)

doc.save(out)
print("Saved:", out)
